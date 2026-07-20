from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass
from io import BytesIO

from app.schemas.documents import (
    EducationAnalysis,
    EvidenceItem,
    ExperienceAnalysis,
    ProjectAnalysis,
    StructuredDocumentAnalysis,
)
from app.services.document_analysis.feedback_generator import generate_category_feedback


@dataclass(frozen=True)
class DocumentProcessingError(Exception):
    code: str
    status_code: int
    message: str


def normalize_text(text: str) -> str:
    text = unicodedata.normalize("NFKC", text).replace("\x00", " ")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def extract_pdf_text(content: bytes, minimum_characters: int) -> str:
    try:
        from pypdf import PdfReader

        text = normalize_text(
            "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(content)).pages)
        )
    except Exception as exc:
        raise DocumentProcessingError("invalid_pdf", 422, "The file is not a readable PDF") from exc
    if not text:
        raise DocumentProcessingError(
            "scanned_pdf", 422, "No selectable text was found; this PDF requires OCR"
        )
    if len(text) < minimum_characters:
        raise DocumentProcessingError(
            "insufficient_text", 422, "The extracted text is too short to analyze"
        )
    return text


def extract_docx_text(content: bytes, minimum_characters: int) -> str:
    try:
        from docx import Document

        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = [
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        ]
        text = normalize_text("\n".join([*paragraphs, *table_cells]))
    except Exception as exc:
        raise DocumentProcessingError(
            "invalid_docx", 422, "The file is not a readable DOCX"
        ) from exc
    if len(text) < minimum_characters:
        raise DocumentProcessingError(
            "insufficient_text", 422, "The extracted text is too short to analyze"
        )
    return text


def extract_txt_text(content: bytes, minimum_characters: int) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp949"):
        try:
            text = normalize_text(content.decode(encoding))
            break
        except UnicodeDecodeError:
            continue
    else:
        raise DocumentProcessingError(
            "invalid_text", 422, "The text file encoding is not supported"
        )
    if len(text) < minimum_characters:
        raise DocumentProcessingError(
            "insufficient_text", 422, "The extracted text is too short to analyze"
        )
    return text


SECTION_ALIASES = {
    "education": ("education", "academic background", "학력", "교육"),
    "work": (
        "experience",
        "work experience",
        "professional experience",
        "employment",
        "internships",
        "경력",
        "업무 경력",
        "인턴십",
    ),
    "research": (
        "research experience",
        "research",
        "연구 경험",
        "연구",
    ),
    "projects": ("projects", "project experience", "selected projects", "프로젝트"),
    "skills": ("skills", "technical skills", "technologies", "기술", "보유 기술"),
    "interests": ("research interests", "interests", "areas of interest", "연구 관심 분야"),
    "campus": (
        "campus & community involvement",
        "campus and community involvement",
        "campus involvement",
        "community involvement",
        "leadership & activities",
        "leadership and activities",
        "extracurricular activities",
        "volunteer experience",
        "awards and activities",
        "activities",
        "교내외 활동",
        "대외 활동",
        "봉사 활동",
    ),
    "publications": ("publications", "publication", "publications and presentations", "논문"),
    "awards": ("awards", "honors", "awards and honors", "수상", "수상 경력"),
    "profile": ("profile", "summary", "objective", "professional summary", "소개"),
    "additional": (
        "additional information",
        "certifications",
        "languages",
        "기타",
        "자격증",
        "언어",
    ),
}

TERM_ALIASES = {
    "Artificial Intelligence": ("artificial intelligence", "인공지능"),
    "Machine Learning": ("machine learning", "머신러닝"),
    "Deep Learning": ("deep learning", "딥러닝"),
    "Computer Vision": ("computer vision", "image recognition", "컴퓨터 비전", "영상 인식"),
    "Natural Language Processing": ("natural language processing", "nlp", "자연어 처리"),
    "Large Language Models": ("large language model", "large language models", "llm", "llms"),
    "Generative AI": ("generative ai", "생성형 ai", "생성형 인공지능"),
    "Multimodal Learning": ("multimodal", "multi-modal", "멀티모달"),
    "Robotics": ("robotics", "robot", "로보틱스", "로봇"),
    "Human-Computer Interaction": ("human-computer interaction", "hci", "인간 컴퓨터 상호작용"),
    "Reinforcement Learning": ("reinforcement learning", "강화학습", "강화 학습"),
    "Data Science": ("data science", "데이터 과학"),
    "Bioinformatics": ("bioinformatics", "생물정보학"),
    "Signal Processing": ("signal processing", "신호 처리"),
    "Distributed Systems": ("distributed systems", "분산 시스템"),
    "Cybersecurity": ("cybersecurity", "computer security", "보안"),
    "Graph Neural Networks": ("graph neural network", "graph neural networks", "gnn"),
    "Diffusion Models": ("diffusion model", "diffusion models", "확산 모델"),
    "Transformers": ("transformer", "transformers"),
    "Object Detection": ("object detection", "객체 탐지"),
    "Image Segmentation": ("image segmentation", "semantic segmentation", "이미지 분할"),
    "Medical Imaging": ("medical imaging", "의료 영상"),
    "Time Series": ("time series", "시계열"),
    "Optimization": ("optimization", "최적화"),
    "Embodied AI": ("embodied ai", "embodied intelligence", "체화 인공지능"),
    "Robot Learning": ("robot learning", "로봇 학습"),
    "3D Vision": ("3d vision", "3d reconstruction", "3차원 비전", "3차원 복원"),
    "Human-Robot Interaction": ("human-robot interaction", "hri", "인간 로봇 상호작용"),
    "Vision-Language Models": ("vision-language model", "vision language model", "vlm"),
    "Recommendation Systems": ("recommendation system", "recommender system", "추천 시스템"),
    "Trajectory Generation": ("trajectory generation", "trajectory processing", "궤적 생성"),
    "Depth Sensing": ("depth sensing", "rgb-d", "depth camera", "깊이 센싱"),
    "Python": ("python",),
    "PyTorch": ("pytorch",),
    "TensorFlow": ("tensorflow",),
    "scikit-learn": ("scikit-learn", "sklearn"),
    "OpenCV": ("opencv",),
    "SQL": ("sql",),
    "Docker": ("docker",),
    "Kubernetes": ("kubernetes",),
    "C++": ("c++",),
    "Java": ("java",),
    "R": ("r language", "r programming"),
    "JavaScript": ("javascript",),
    "TypeScript": ("typescript",),
    "Bash": ("bash",),
    "FastAPI": ("fastapi",),
    "React": ("react",),
    "PostgreSQL": ("postgresql", "postgres"),
    "MySQL": ("mysql",),
    "SQLAlchemy": ("sqlalchemy",),
    "Pydantic": ("pydantic",),
    "NumPy": ("numpy",),
    "Pandas": ("pandas",),
    "Git": ("git",),
    "GitHub": ("github",),
    "Linux": ("linux", "wsl"),
    "Jupyter": ("jupyter",),
    "CMake": ("cmake",),
    "ROS": ("ros", "robot operating system"),
    "CUDA": ("cuda",),
    "AWS": ("aws", "amazon web services"),
}

DOMAIN_TERMS = {
    "Artificial Intelligence",
    "Machine Learning",
    "Deep Learning",
    "Computer Vision",
    "Natural Language Processing",
    "Large Language Models",
    "Generative AI",
    "Multimodal Learning",
    "Robotics",
    "Human-Computer Interaction",
    "Reinforcement Learning",
    "Data Science",
    "Bioinformatics",
    "Signal Processing",
    "Distributed Systems",
    "Cybersecurity",
    "Graph Neural Networks",
    "Diffusion Models",
    "Transformers",
    "Object Detection",
    "Image Segmentation",
    "Medical Imaging",
    "Time Series",
    "Optimization",
    "Embodied AI",
    "Robot Learning",
    "3D Vision",
    "Human-Robot Interaction",
    "Vision-Language Models",
    "Recommendation Systems",
    "Trajectory Generation",
    "Depth Sensing",
}
SKILL_TERMS = set(TERM_ALIASES) - DOMAIN_TERMS
DATE_PATTERN = re.compile(
    r"(?P<start>(?:19|20)\d{2}(?:[./-]\d{1,2})?)\s*(?:[-–—~]|to)\s*"
    r"(?P<end>(?:(?:19|20)\d{2}(?:[./-]\d{1,2})?|present|current|현재))",
    re.IGNORECASE,
)


def _heading_key(line: str) -> str | None:
    candidate = re.sub(r"[:|]$", "", line.strip()).casefold()
    candidate = re.sub(r"\s+", " ", candidate)
    for key, aliases in SECTION_ALIASES.items():
        if candidate in aliases:
            return key
    return None


def _sections(lines: list[str]) -> tuple[dict[str, list[str]], list[str]]:
    result = {key: [] for key in SECTION_ALIASES}
    unsectioned: list[str] = []
    current: str | None = None
    for line in lines:
        heading = _heading_key(line)
        if heading:
            current = heading
        elif current:
            result[current].append(line)
        else:
            unsectioned.append(line)
    return result, unsectioned


def _clean_bullet(line: str) -> str:
    return re.sub(r"^(?:[•●▪◦\uf0b7*]|[-–—]\s+)\s*", "", line).strip()


def _date_range(text: str) -> tuple[str, str]:
    match = DATE_PATTERN.search(text)
    if match:
        return match.group("start"), match.group("end")
    years = re.findall(r"(?:19|20)\d{2}(?:[./-]\d{1,2})?", text)
    return (years[0], years[-1] if len(years) > 1 else "") if years else ("", "")


def _is_detail(line: str) -> bool:
    return bool(re.match(r"^(?:[•●▪◦\uf0b7*]|[-–—]\s+)", line)) or len(line) > 115


def _group_entries(lines: list[str]) -> list[tuple[str, list[str]]]:
    entries: list[tuple[str, list[str]]] = []
    title = ""
    details: list[str] = []
    for raw in lines:
        line = _clean_bullet(raw)
        if not line:
            continue
        starts_entry = "|" in raw and not _is_detail(raw)
        if title and not starts_entry:
            details.append(line)
            continue
        if title:
            entries.append((title, details))
        title, details = line, []
    if title:
        entries.append((title, details))
    return entries


def _split_header(header: str) -> tuple[str, str, str, str, str]:
    start, end = _date_range(header)
    without_dates = DATE_PATTERN.sub("", header)
    parts = [part.strip(" ,|-–—") for part in re.split(r"\s*[|·]\s*", without_dates)]
    parts = [part for part in parts if part]
    title = parts[0] if parts else without_dates.strip()
    organization = parts[1] if len(parts) > 1 else ""
    location = parts[2] if len(parts) > 2 else ""
    return title[:300], organization[:300], location[:200], start, end


def _experience(lines: list[str]) -> list[ExperienceAnalysis]:
    items: list[ExperienceAnalysis] = []
    for header, details in _group_entries(lines)[:10]:
        title, organization, location, start, end = _split_header(header)
        combined = " ".join([header, *details])
        if not start:
            start, end = _date_range(combined)
        if not location:
            location = next(
                (
                    detail
                    for detail in details
                    if re.search(
                        r"korea|seoul|pohang|city|province|대한민국|서울|포항", detail, re.I
                    )
                ),
                "",
            )[:200]
        visible_details = [
            detail
            for detail in details
            if detail
            and detail != location
            and not DATE_PATTERN.fullmatch(detail)
            and not re.fullmatch(r"(?:19|20)\d{2}(?:[./-]\d{1,2})?", detail)
        ]
        items.append(
            ExperienceAnalysis(
                title=title,
                organization=organization,
                location=location,
                start_date=start,
                end_date=end,
                details=visible_details[:8],
            )
        )
    return items


def _education(lines: list[str], unsectioned: list[str]) -> list[EducationAnalysis]:
    candidates = lines or [
        line
        for line in unsectioned
        if re.search(
            r"\b(?:b\.?s\.?|m\.?s\.?|ph\.?d|bachelor|master|university|college)\b|대학교|대학원|학사|석사|박사",
            line,
            re.IGNORECASE,
        )
    ]
    degree_pattern = re.compile(
        r"bachelor|master|ph\.?d|b\.?s\.?|m\.?s\.?|학사|석사|박사", re.IGNORECASE
    )
    grouped: list[tuple[str, list[str]]] = []
    header = ""
    details: list[str] = []
    for line in candidates:
        if degree_pattern.search(line) and (not header or "|" in line):
            if header:
                grouped.append((header, details))
            header, details = line, []
        elif header:
            details.append(_clean_bullet(line))
    if header:
        grouped.append((header, details))

    items: list[EducationAnalysis] = []
    for header, details in grouped[:6]:
        combined = " ".join([header, *details])
        start, end = _date_range(combined)
        parts = [part.strip() for part in re.split(r"\s*[|·]\s*", header) if part.strip()]
        degree_index = next(
            (index for index, part in enumerate(parts) if degree_pattern.search(part)),
            None,
        )
        institution_index = next(
            (
                index
                for index, part in enumerate(parts)
                if re.search(r"university|college|institute|대학교|대학원", part, re.I)
            ),
            None,
        )
        items.append(
            EducationAnalysis(
                degree=parts[degree_index] if degree_index is not None else header[:300],
                institution=parts[institution_index] if institution_index is not None else "",
                location=next(
                    (
                        detail
                        for detail in details
                        if re.search(
                            r"korea|seoul|pohang|city|province|대한민국|서울|포항", detail, re.I
                        )
                    ),
                    "",
                )[:200],
                start_date=start,
                end_date=end,
                details=[
                    detail
                    for detail in details
                    if detail and detail not in {start, end} and not DATE_PATTERN.fullmatch(detail)
                ][:6],
            )
        )
    return items


def _matches(text: str, alias: str) -> list[re.Match[str]]:
    escaped = re.escape(alias).replace(r"\ ", r"[\s_-]+")
    if re.fullmatch(r"[a-z0-9 +.#-]+", alias):
        escaped = rf"(?<![\w]){escaped}(?![\w])"
    return list(re.finditer(escaped, text, re.IGNORECASE))


def _extract_terms(
    text: str, priority_text: str
) -> tuple[list[str], dict[str, float], dict[str, list[EvidenceItem]]]:
    hits: dict[str, list[re.Match[str]]] = {}
    priority_hits: dict[str, int] = {}
    for label, aliases in TERM_ALIASES.items():
        matches = [match for alias in aliases for match in _matches(text, alias)]
        if matches:
            hits[label] = matches
            priority_hits[label] = sum(len(_matches(priority_text, alias)) for alias in aliases)
    ordered = sorted(
        hits,
        key=lambda label: (
            -priority_hits[label],
            -len(hits[label]),
            label.casefold(),
        ),
    )
    weights = {
        label: min(
            1.0,
            round(
                0.45 + 0.08 * len(hits[label]) + min(0.3, 0.15 * priority_hits[label]),
                2,
            ),
        )
        for label in ordered[:30]
    }
    lines = text.splitlines()
    evidence = {
        "keywords": [
            EvidenceItem(
                value=label,
                confidence=weights[label],
                evidence=next(
                    (line[:280] for line in lines if _matches(line, TERM_ALIASES[label][0])), label
                ),
            )
            for label in ordered[:30]
        ]
    }
    return ordered[:30], weights, evidence


class LocalRuleBasedCvAnalyzer:
    def analyze(self, text: str) -> StructuredDocumentAnalysis:
        text = normalize_text(text)
        lines = text.splitlines()
        sections, unsectioned = _sections(lines)
        education = _education(sections["education"], unsectioned)
        work = _experience(sections["work"])
        research = _experience(sections["research"])
        campus = _experience(sections["campus"])
        priority_text = " ".join(
            sections["skills"] + sections["interests"] + sections["research"] + sections["projects"]
        )
        keywords, keyword_weights, evidence = _extract_terms(text, priority_text)
        skills = [term for term in keywords if term in SKILL_TERMS]
        interests = [term for term in keywords if term in DOMAIN_TERMS]
        projects = []
        for header, details in _group_entries(sections["projects"])[:10]:
            title, organization, location, start, end = _split_header(header)
            project_text = " ".join([header, *details])
            if not start:
                start, end = _date_range(project_text)
            visible_details = [
                detail
                for detail in details
                if detail
                and not DATE_PATTERN.fullmatch(detail)
                and not re.fullmatch(r"(?:19|20)\d{2}(?:[./-]\d{1,2})?", detail)
            ]
            technologies = [
                term
                for term in keywords
                if term in SKILL_TERMS
                and any(_matches(project_text, alias) for alias in TERM_ALIASES[term])
            ]
            projects.append(
                ProjectAnalysis(
                    name=title or "Untitled project",
                    organization=organization,
                    location=location,
                    start_date=start,
                    end_date=end,
                    description=" ".join(visible_details[:3])[:2000],
                    details=visible_details[:8],
                    technologies=technologies,
                )
            )
        strengths = [
            f"Repeated evidence of {term}" for term in keywords if keyword_weights[term] >= 0.75
        ][:5]
        missing = [
            label
            for label, value in (
                ("Education", education),
                ("Research experience", research),
                ("Projects", projects),
                ("Technical skills", skills),
            )
            if not value
        ]
        evidence["skills"] = [item for item in evidence["keywords"] if item.value in skills]
        evidence["research_interests"] = [
            item for item in evidence["keywords"] if item.value in interests
        ]
        feedback = generate_category_feedback(
            text,
            education=education,
            skills=skills,
            projects_count=len(projects),
            research_experience=research,
            work_experience=work,
            campus_community_involvement=campus,
        )
        summary_parts = []
        if interests:
            summary_parts.append(f"Research focus: {', '.join(interests[:4])}")
        if skills:
            summary_parts.append(f"Core skills: {', '.join(skills[:5])}")
        short_summary = (
            ". ".join(summary_parts)
            or "CV sections were organized, but no supported research keywords were found"
        )
        return StructuredDocumentAnalysis(
            education=education,
            work_experience=work,
            campus_community_involvement=campus,
            research_experience=research,
            projects=projects,
            skills=skills,
            research_interests=interests,
            strengths=strengths,
            missing_information=missing,
            keywords=keywords,
            keyword_weights=keyword_weights,
            short_summary=f"{short_summary}.",
            evidence_items=evidence,
            category_feedback=feedback,
        )


def analyze_document_text(text: str, **_: object) -> StructuredDocumentAnalysis:
    return LocalRuleBasedCvAnalyzer().analyze(text)
